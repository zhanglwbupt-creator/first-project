from docx import Document
import json

doc = Document('01_产品需求文档PRD.docx')

# 提取所有段落
print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'None'
        print(f"[{style}] {para.text}")

# 提取所有表格
print("\n\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {t_idx + 1} ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"Row {r_idx}: {cells}")
