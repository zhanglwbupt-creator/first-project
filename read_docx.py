from docx import Document
import sys

def read_docx(filepath):
    try:
        doc = Document(filepath)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{para.text}")
        print("\n=== TABLES ===")
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                print(" | ".join(cells))
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
    else:
        print("Usage: python read_docx.py <file.docx>")
